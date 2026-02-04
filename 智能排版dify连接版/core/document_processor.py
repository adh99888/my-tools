"""
文档处理器模块
负责文档的加载、处理和保存
"""

import os
import re
import logging
from pathlib import Path
from typing import Optional, Dict, Any, Tuple
from datetime import datetime
import threading
import queue

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import PyPDF2

# 导入后处理器
from .long_document_processor import DocumentPostProcessor
# 导入提示词配置加载函数
from .template_manager import load_prompt_config

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """文档处理器类"""
    
    def __init__(self, config_manager, model_manager, template_manager, api_client):
        """
        初始化文档处理器
        
        Args:
            config_manager: 配置管理器实例
            model_manager: 模型管理器实例
            template_manager: 模板管理器实例
            api_client: API客户端实例
        """
        self.config_manager = config_manager
        self.model_manager = model_manager
        self.template_manager = template_manager
        self.api_client = api_client
        
        # 当前处理的文档信息
        self.current_file = ""
        self.original_title = ""
        self.current_content = ""
        self.processed_content = ""
        
        # 处理状态
        self.is_processing = False
        
        # 消息队列用于线程间通信
        self.message_queue = queue.Queue()
    
    def load_file(self, file_path: str) -> Tuple[bool, str]:
        """
        加载文档文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            (success: bool, content: str) 加载是否成功和内容
        """
        self.current_file = file_path
        
        try:
            # 根据文件类型读取
            if file_path.lower().endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif file_path.lower().endswith('.docx'):
                from docx import Document
                doc = Document(file_path)
                content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            elif file_path.lower().endswith('.pdf'):
                content = self._read_pdf(file_path)
            else:
                return False, f"不支持的文件格式: {file_path}"
            
            self.current_content = content
            
            # 提取标题
            self.original_title = self.extract_title(content)
            
            logger.info(f"文档加载成功: {os.path.basename(file_path)}")
            return True, content
            
        except Exception as e:
            logger.error(f"加载文件失败: {str(e)}")
            return False, f"加载失败: {str(e)}"
    
    def _read_pdf(self, filename: str) -> str:
        """读取PDF文件"""
        content = ""
        try:
            with open(filename, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        content += page_text + "\n\n"
            
            return content
        except Exception as e:
            raise Exception(f"PDF读取失败: {str(e)}")
    
    def extract_title(self, content: str, max_lines: int = 5) -> str:
        """
        从内容中提取标题
        
        Args:
            content: 文档内容
            max_lines: 最多检查的行数
            
        Returns:
            提取的标题
        """
        lines = content.strip().split('\n')
        for line in lines[:max_lines]:
            line = line.strip()
            if line and len(line) < 100 and not line.startswith((' ', '\t', '#', '-', '*')):
                return line
        return "文档标题"
    
    def _enhance_content_with_instructions(self, content: str, title: str) -> str:
        """
        在内容前添加排版指令，优先使用配置文件中的提示词
        
        Args:
            content: 原始文档内容
            title: 文档标题
            
        Returns:
            增强后的内容
        """
        # 加载提示词配置
        prompt_config = load_prompt_config()
        
        # 获取当前模板ID
        template_id = self.template_manager.current_template
        logger.info(f"当前模板ID: {template_id}")
        
        # 确定使用的提示词
        final_prompt = ""
        
        # 检查开关状态
        if prompt_config.get("enabled", False):
            templates_config = prompt_config.get("templates", {})
            
            # 优先使用模板专用提示词
            if template_id in templates_config:
                custom_prompt = templates_config[template_id].get("prompt", "")
                if custom_prompt:
                    final_prompt = custom_prompt
                    logger.info(f"使用自定义提示词 - 模板: {template_id}")
                else:
                    final_prompt = prompt_config.get("default_prompt", "")
                    logger.info(f"模板未配置提示词，使用默认提示词 - 模板: {template_id}")
            else:
                final_prompt = prompt_config.get("default_prompt", "")
                logger.info(f"模板未在配置中，使用默认提示词")
        else:
            # 使用原提示词（硬编码的）
            original_prompt = """【排版要求】（请严格执行并确保返回的格式符合Word文档要求）：

1. 标题处理：
   - 使用#表示一级标题，##表示二级标题，###表示三级标题
   - 一级标题只需在开头出现一次，不要重复
   - 中文小标题（如一、二、三）请转换为##二级标题格式

2. 段落优化：
   - 段落之间保留一个空行作为分隔
   - 不要使用---、===、――等分隔线，Word会自动处理分隔
   - 保持段落自然流畅

3. 格式清理：
   - 去除所有孤立的符号：•、*、-、#等（除非用于列表）
   - 列表使用-或*开头，后面跟一个空格
   - 规范标点使用，中文使用全角标点

4. 特别强调：
   - 不要添加任何装饰性符号（如■、◆、☆等）
   - 不要添加虚线或重复的分隔线
   - 确保返回的格式可以直接导入Word

【原文标题】：{title}

【原文内容】开始：
"""
            final_prompt = original_prompt.format(title=title)
            logger.info("提示词优化功能已禁用，使用原提示词")
        
        # 构建最终内容
        return final_prompt + content
    
    def _post_process_content(self, content: str) -> str:
        """
        后处理：清理AI返回结果中的问题
        
        Args:
            content: AI处理后的内容
            
        Returns:
            清理后的内容
        """
        if not content:
            return content
            
        # 1. 去除排版指令部分（如果AI返回时包含了）
        instruction_markers = ["【排版要求】", "【特别强调】", "【原文标题】", "【原文内容】开始："]
        for marker in instruction_markers:
            if marker in content:
                # 只保留指令之后的内容
                parts = content.split(marker)
                if len(parts) > 1:
                    # 取最后一部分作为实际内容
                    content = parts[-1]
        
        # 2. 清理多余的分隔线和装饰符号
        lines = content.split('\n')
        processed_lines = []
        
        for line in lines:
            stripped = line.strip()
            
            # 跳过空行（空行在下一步处理）
            if not stripped:
                processed_lines.append(line)
                continue
            
            # 3. 清理各种分隔线和装饰符号
            # 严格跳过所有形式的分隔线
            if re.match(r'^[―\-—=_*#\.~·]{4,}$', stripped):
                logger.debug(f"跳过分隔线: {stripped}")
                continue
            
            # 4. 清理"生成时间"行（我们会在Word中单独添加）
            if stripped.startswith('生成时间:') and ':' in stripped:
                logger.debug(f"跳过生成时间行: {stripped}")
                continue
            
            # 5. 清理行首的孤立符号
            # 保留用于列表的符号（- 或 * 后跟空格）
            if re.match(r'^[-*•]\s', stripped):
                # 这是列表项，保留
                processed_lines.append(line)
            elif re.match(r'^[•*\-—=_#\.~·]{1,3}$', stripped):
                # 孤立的符号，跳过
                logger.debug(f"跳过孤立符号: {stripped}")
                continue
            else:
                # 普通文本行
                processed_lines.append(line)
        
        # 6. 合并并处理空行
        content = '\n'.join(processed_lines)
        
        # 7. 去除多余的空行（最多保留一个空行）
        lines = content.split('\n')
        processed_lines = []
        prev_empty = False
        
        for line in lines:
            stripped = line.strip()
            if not stripped:  # 空行
                if not prev_empty:
                    processed_lines.append('')
                    prev_empty = True
            else:
                processed_lines.append(line)
                prev_empty = False
        
        # 8. 确保文档开头不是空行
        while processed_lines and not processed_lines[0].strip():
            processed_lines.pop(0)
        
        # 9. 确保文档结尾不是空行
        while processed_lines and not processed_lines[-1].strip():
            processed_lines.pop()
        
        # 10. 合并结果
        result = '\n'.join(processed_lines)
        
        # 11. 如果有DocumentPostProcessor，也使用它的处理
        try:
            from .long_document_processor import DocumentPostProcessor
            result = DocumentPostProcessor.process(result)
        except Exception as e:
            logger.debug(f"DocumentPostProcessor处理跳过: {str(e)}")
        
        return result
    
    def process_document(self, content: str, model_id: str = None) -> Tuple[bool, str]:
        """
        处理文档内容（增强版）
        
        Args:
            content: 文档内容
            model_id: 模型ID，如果为None则自动选择
            
        Returns:
            (success: bool, result: str) 处理是否成功和处理结果
        """
        if not content:
            return False, "文档内容为空"
        
        # 提取标题用于增强内容
        title = self.extract_title(content)
        
        # 增强内容，添加排版指令
        enhanced_content = self._enhance_content_with_instructions(content, title)
        
        # 检查内容长度
        char_count = len(content)
        
        # 自动选择模型（如果需要）
        if model_id is None:
            model_id = self.model_manager.get_model_for_content(content)
        
        model_config = self.model_manager.get_model_config(model_id)
        if not model_config:
            return False, f"模型 '{model_id}' 不存在"
        
        # 检查token限制
        max_tokens = model_config.get('max_tokens', 8192)
        estimated_tokens = char_count // 3  # 粗略估计
        
        if estimated_tokens > max_tokens * 0.7:  # 使用70%的安全余量
            logger.warning(
                f"文档较长 ({char_count}字符，约{estimated_tokens}tokens)，"
                f"当前模型最大支持{max_tokens}tokens"
            )
        
        # 调用API处理 - 使用增强后的内容
        success, result = self.api_client.process_with_retry(
            enhanced_content, 
            self.config_manager.api_config.get('max_retries', 3)
        )
        
        if not success:
            return False, result
        
        # 应用后处理优化
        try:
            # 首先使用DocumentPostProcessor处理
            result = DocumentPostProcessor.process(result)
        except Exception as e:
            logger.warning(f"DocumentPostProcessor处理失败: {str(e)}")
        
        # 然后应用我们的后处理
        try:
            result = self._post_process_content(result)
        except Exception as e:
            logger.warning(f"后处理失败，但继续使用原始结果: {str(e)}")
        
        self.processed_content = result
        return True, result
    
    def save_as_word(self, content: str, title: str, 
                    template_name: str = None, output_path: str = None) -> Tuple[bool, str]:
        """
        保存为Word文档（修复标题重复和样式问题）
        
        Args:
            content: 要保存的内容
            title: 文档标题
            template_name: 模板名称
            output_path: 输出路径（如果提供，则使用此路径）
            
        Returns:
            (success: bool, file_path: str) 保存是否成功和文件路径
        """
        try:
            # 获取模板
            if template_name is None:
                template_name = self.template_manager.current_template
            
            template_config = self.template_manager.get_template(template_name)
            if not template_config:
                return False, f"模板 '{template_name}' 不存在"
            
            # 生成文件名
            safe_title = re.sub(r'[\\/*?:"<>|]', '_', title)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            default_name = f"{safe_title}_{timestamp}.docx"
            
            # 确定最终文件路径
            if output_path:
                # 如果output_path是目录，则拼接文件名
                if os.path.isdir(output_path):
                    file_path = Path(output_path) / default_name
                else:
                    # 如果output_path是文件路径，直接使用
                    file_path = Path(output_path)
                    # 确保文件扩展名正确
                    if file_path.suffix.lower() != '.docx':
                        file_path = file_path.with_suffix('.docx')
            else:
                # 默认保存到当前工作目录
                file_path = Path.cwd() / default_name
            
            # 确保目录存在
            file_path.parent.mkdir(parents=True, exist_ok=True)
            
            # 创建文档
            doc = Document()
            
            # 应用模板样式
            if not self.template_manager.apply_document_styles(doc, template_config):
                logger.warning("样式应用部分失败，使用默认样式")
            
            # 添加标题（确保居中）
            title_text = title or self.original_title or "文档标题"
            
            # 检查模板中是否有标题样式配置
            if 'title' in template_config:
                title_style = template_config['title']
                title_para = doc.add_heading(title_text, level=0)
                
                # 应用标题样式
                if 'font_size' in title_style:
                    title_para.style.font.size = Pt(title_style['font_size'])
                if 'alignment' in title_style and title_style['alignment'] == 'center':
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # 默认居中
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # 默认添加居中标题
                title_para = doc.add_heading(title_text, level=0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加元信息
            meta = doc.add_paragraph()
            meta_run = meta.add_run(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加简单的分隔线（一条即可）
            separator = doc.add_paragraph('―' * 40)
            separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 处理内容，修复标题重复问题
            lines = content.split('\n')
            skip_main_title = True  # 标记是否跳过了主标题
            prev_was_paragraph = False
            
            # 处理中文数字标题映射
            chinese_numbers = {
                '一': '1', '二': '2', '三': '3', '四': '4', '五': '5',
                '六': '6', '七': '7', '八': '8', '九': '9', '十': '10'
            }
            
            for line in lines:
                line = line.strip()
                if not line:
                    if not prev_was_paragraph:  # 避免连续空行
                        doc.add_paragraph()
                        prev_was_paragraph = True
                    continue
                
                # 跳过任何看起来像分隔线的行
                if re.match(r'^[―\-—=_*#\.~·]{4,}$', line):
                    logger.debug(f"跳过内容中的分隔线: {line}")
                    continue
                
                # 跳过"生成时间"行（我们已经添加过了）
                if line.startswith('生成时间:'):
                    logger.debug(f"跳过内容中的生成时间行: {line}")
                    continue
                
                # 检查是否为标题行
                if line.startswith('# '):
                    # 一级标题：跳过，因为已经添加过了
                    if skip_main_title:
                        skip_main_title = False
                        continue
                    else:
                        # 如果还有一级标题，转为二级标题（避免重复）
                        heading = doc.add_heading(line[2:], level=2)
                        prev_was_paragraph = False
                elif line.startswith('## '):
                    heading = doc.add_heading(line[3:], level=2)
                    prev_was_paragraph = False
                elif line.startswith('### '):
                    heading = doc.add_heading(line[4:], level=3)
                    prev_was_paragraph = False
                # 处理中文数字标题（如一、二、三）
                elif re.match(r'^[一二三四五六七八九十]、\s*', line):
                    # 移除中文数字和顿号，添加为二级标题
                    clean_line = re.sub(r'^[一二三四五六七八九十]、\s*', '', line)
                    heading = doc.add_heading(clean_line, level=2)
                    prev_was_paragraph = False
                # 处理带括号的中文标题（如（一）、（二））
                elif re.match(r'^（[一二三四五六七八九十]）\s*', line):
                    clean_line = re.sub(r'^（[一二三四五六七八九十]）\s*', '', line)
                    heading = doc.add_heading(clean_line, level=3)
                    prev_was_paragraph = False
                # 处理数字标题（如1.、2.）
                elif re.match(r'^\d+\.\s*', line):
                    # 判断是二级还是三级标题
                    match = re.match(r'^(\d+)\.\s*(.+)', line)
                    if match:
                        num = int(match.group(1))
                        text = match.group(2)
                        if num <= 10:  # 假设10以内的是二级标题
                            heading = doc.add_heading(text, level=2)
                        else:
                            heading = doc.add_heading(text, level=3)
                        prev_was_paragraph = False
                elif line.startswith('- ') or line.startswith('* '):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(line[2:])
                    prev_was_paragraph = True
                elif re.match(r'^\d+[\.\)] ', line):
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(re.sub(r'^\d+[\.\)] ', '', line))
                    prev_was_paragraph = True
                else:
                    doc.add_paragraph(line)
                    prev_was_paragraph = True
            
            # 保存文档
            doc.save(str(file_path))
            
            logger.info(f"文档保存成功: {file_path}")
            return True, str(file_path)
            
        except Exception as e:
            logger.error(f"保存文档失败: {str(e)}")
            return False, f"保存失败: {str(e)}"
    
    def get_stats(self, original_content: str, processed_content: str) -> Dict[str, Any]:
        """
        获取处理统计信息
        
        Args:
            original_content: 原文内容
            processed_content: 处理后内容
            
        Returns:
            统计信息字典
        """
        original_length = len(original_content)
        processed_length = len(processed_content)
        
        if original_length > 0:
            change_rate = (processed_length - original_length) / original_length * 100
        else:
            change_rate = 0
        
        return {
            'original_length': original_length,
            'processed_length': processed_length,
            'change_rate': change_rate,
            'characters_added': processed_length - original_length
        }
    
    def start_processing_thread(self, content: str, model_id: str = None, 
                               callback=None, error_callback=None):
        """
        在单独的线程中启动文档处理
        
        Args:
            content: 文档内容
            model_id: 模型ID
            callback: 处理成功的回调函数
            error_callback: 处理失败的回调函数
        """
        if self.is_processing:
            if error_callback:
                error_callback("正在处理中，请稍候...")
            return
        
        self.is_processing = True
        
        def process_thread():
            try:
                success, result = self.process_document(content, model_id)
                self.is_processing = False
                
                if success and callback:
                    callback(result)
                elif not success and error_callback:
                    error_callback(result)
            except Exception as e:
                self.is_processing = False
                if error_callback:
                    error_callback(f"处理过程中发生错误: {str(e)}")
        
        thread = threading.Thread(target=process_thread)
        thread.daemon = True
        thread.start()
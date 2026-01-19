"""
文档处理器模块
负责文档的加载、处理和保存
"""

import os
import re
import logging
from pathlib import Path
from typing import Optional, Dict, Any, Tuple
from datetime import datetime
import threading
import queue

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2

# 导入后处理器
from .long_document_processor import DocumentPostProcessor

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """文档处理器类"""
    
    def __init__(self, config_manager, model_manager, template_manager, api_client):
        """
        初始化文档处理器
        
        Args:
            config_manager: 配置管理器实例
            model_manager: 模型管理器实例
            template_manager: 模板管理器实例
            api_client: API客户端实例
        """
        self.config_manager = config_manager
        self.model_manager = model_manager
        self.template_manager = template_manager
        self.api_client = api_client
        
        # 当前处理的文档信息
        self.current_file = ""
        self.original_title = ""
        self.current_content = ""
        self.processed_content = ""
        
        # 处理状态
        self.is_processing = False
        
        # 消息队列用于线程间通信
        self.message_queue = queue.Queue()
    
    def load_file(self, file_path: str) -> Tuple[bool, str]:
        """
        加载文档文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            (success: bool, content: str) 加载是否成功和内容
        """
        self.current_file = file_path
        
        try:
            # 根据文件类型读取
            if file_path.lower().endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif file_path.lower().endswith('.docx'):
                from docx import Document
                doc = Document(file_path)
                content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            elif file_path.lower().endswith('.pdf'):
                content = self._read_pdf(file_path)
            else:
                return False, f"不支持的文件格式: {file_path}"
            
            self.current_content = content
            
            # 提取标题
            self.original_title = self.extract_title(content)
            
            logger.info(f"文档加载成功: {os.path.basename(file_path)}")
            return True, content
            
        except Exception as e:
            logger.error(f"加载文件失败: {str(e)}")
            return False, f"加载失败: {str(e)}"
    
    def _read_pdf(self, filename: str) -> str:
        """读取PDF文件"""
        content = ""
        try:
            with open(filename, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        content += page_text + "\n\n"
                
            return content
        except Exception as e:
            raise Exception(f"PDF读取失败: {str(e)}")
    
    def extract_title(self, content: str, max_lines: int = 5) -> str:
        """
        从内容中提取标题
        
        Args:
            content: 文档内容
            max_lines: 最多检查的行数
            
        Returns:
            提取的标题
        """
        lines = content.strip().split('\n')
        for line in lines[:max_lines]:
            line = line.strip()
            if line and len(line) < 100 and not line.startswith((' ', '\t', '#', '-', '*')):
                return line
        return "文档标题"
    
    def process_document(self, content: str, model_id: str = None) -> Tuple[bool, str]:
        """
        处理文档内容
        
        Args:
            content: 文档内容
            model_id: 模型ID，如果为None则自动选择
            
        Returns:
            (success: bool, result: str) 处理是否成功和处理结果
        """
        if not content:
            return False, "文档内容为空"
        
        # 检查内容长度
        char_count = len(content)
        
        # 自动选择模型（如果需要）
        if model_id is None:
            model_id = self.model_manager.get_model_for_content(content)
        
        model_config = self.model_manager.get_model_config(model_id)
        if not model_config:
            return False, f"模型 '{model_id}' 不存在"
        
        # 检查token限制
        max_tokens = model_config.get('max_tokens', 8192)
        estimated_tokens = char_count // 3  # 粗略估计
        
        if estimated_tokens > max_tokens * 0.7:  # 使用70%的安全余量
            logger.warning(
                f"文档较长 ({char_count}字符，约{estimated_tokens}tokens)，"
                f"当前模型最大支持{max_tokens}tokens"
            )
        
        # 调用API处理
        success, result = self.api_client.process_with_retry(
            content, 
            self.config_manager.api_config.get('max_retries', 3)
        )
        
        if success:
            # 集成点1：后处理优化
            # 应用文档后处理优化
            try:
                result = DocumentPostProcessor.process(result)
            except Exception as e:
                logger.warning(f"后处理优化失败，但继续使用原始结果: {str(e)}")
            
            self.processed_content = result
            return True, result
        else:
            return False, result
    
    def save_as_word(self, content: str, title: str, 
                    template_name: str = None) -> Tuple[bool, str]:
        """
        保存为Word文档
        
        Args:
            content: 要保存的内容
            title: 文档标题
            template_name: 模板名称
            
        Returns:
            (success: bool, file_path: str) 保存是否成功和文件路径
        """
        try:
            # 获取模板
            if template_name is None:
                template_name = self.template_manager.current_template
            
            template_config = self.template_manager.get_template(template_name)
            if not template_config:
                return False, f"模板 '{template_name}' 不存在"
            
            # 生成文件名
            safe_title = re.sub(r'[\\/*?:"<>|]', '_', title)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            default_name = f"{safe_title}_{timestamp}.docx"
            
            # 创建文档
            doc = Document()
            
            # 应用模板样式
            if not self.template_manager.apply_document_styles(doc, template_config):
                logger.warning("样式应用部分失败，使用默认样式")
            
            # 添加标题
            title_text = title or self.original_title or "文档标题"
            if 'title' in template_config:
                title_style = template_config['title']
                title_para = doc.add_heading(title_text, level=0)
                
                # 应用标题样式
                if 'font_size' in title_style:
                    title_para.style.font.size = Pt(title_style['font_size'])
                if 'alignment' in title_style and title_style['alignment'] == 'center':
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(title_text, level=0)
            
            # 添加元信息
            meta = doc.add_paragraph()
            meta.add_run(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph('―' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加内容
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue
                
                # 识别标题和列表
                # 跳过顶级标题（避免与用户输入的标题重复）
                if line.startswith('# '):
                    # 修改点4：跳过顶级标题行，避免重复
                    # 只添加非标题内容或使用正文格式
                    doc.add_paragraph(line[2:])
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(line[2:])
                elif re.match(r'^\d+[\.\)] ', line):
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(re.sub(r'^\d+[\.\)] ', '', line))
                else:
                    doc.add_paragraph(line)
            
            # 保存文档
            file_path = Path.cwd() / default_name
            doc.save(str(file_path))
            
            logger.info(f"文档保存成功: {file_path}")
            return True, str(file_path)
            
        except Exception as e:
            logger.error(f"保存文档失败: {str(e)}")
            return False, f"保存失败: {str(e)}"
    
    def get_stats(self, original_content: str, processed_content: str) -> Dict[str, Any]:
        """
        获取处理统计信息
        
        Args:
            original_content: 原文内容
            processed_content: 处理后内容
            
        Returns:
            统计信息字典
        """
        original_length = len(original_content)
        processed_length = len(processed_content)
        
        if original_length > 0:
            change_rate = (processed_length - original_length) / original_length * 100
        else:
            change_rate = 0
        
        return {
            'original_length': original_length,
            'processed_length': processed_length,
            'change_rate': change_rate,
            'characters_added': processed_length - original_length
        }
    
    def start_processing_thread(self, content: str, model_id: str = None, 
                               callback=None, error_callback=None):
        """
        在单独的线程中启动文档处理
        
        Args:
            content: 文档内容
            model_id: 模型ID
            callback: 处理成功的回调函数
            error_callback: 处理失败的回调函数
        """
        if self.is_processing:
            if error_callback:
                error_callback("正在处理中，请稍候...")
            return
        
        self.is_processing = True
        
        def process_thread():
            try:
                success, result = self.process_document(content, model_id)
                self.is_processing = False
                
                if success and callback:
                    callback(result)
                elif not success and error_callback:
                    error_callback(result)
            except Exception as e:
                self.is_processing = False
                if error_callback:
                    error_callback(f"处理过程中发生错误: {str(e)}")
        
        thread = threading.Thread(target=process_thread)
        thread.daemon = True
        thread.start()